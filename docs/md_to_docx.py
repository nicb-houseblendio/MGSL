"""Convert architecture_summary.md to .docx"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INPUT = r"D:\HouseBlend\Clients\MGSL\Implementation\docs\architecture_summary.md"
OUTPUT = r"D:\HouseBlend\Clients\MGSL\Implementation\docs\architecture_summary.docx"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x0F, 0x26, 0x41)

def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code = False
code_buf = []
in_table = False
table_headers = []
table_rows = []

while i < len(lines):
    line = lines[i].rstrip('\n')

    if line.startswith('```'):
        if in_code:
            add_code_block('\n'.join(code_buf))
            code_buf = []
            in_code = False
        else:
            if in_table:
                add_table(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if line.startswith('|') and '|' in line[1:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not in_table:
            table_headers = cells
            in_table = True
        elif all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
            pass  # separator row
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.strip() == '---':
        pass
    elif line.strip() == '':
        pass
    elif line.startswith('- **'):
        m = re.match(r'^- \*\*(.+?)\*\*:?\s*(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(m.group(1))
            run.bold = True
            run.font.size = Pt(10)
            if m.group(2):
                rest = p.add_run(': ' + m.group(2) if not m.group(2).startswith(':') else m.group(2))
                rest.font.size = Pt(10)
        else:
            doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif re.match(r'^\d+\.\s', line):
        m = re.match(r'^\d+\.\s(.*)', line)
        text = m.group(1)
        p = doc.add_paragraph(style='List Number')
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
    else:
        text = line
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
            else:
                p.add_run(part)
    i += 1

if in_table:
    add_table(table_headers, table_rows)
if in_code:
    add_code_block('\n'.join(code_buf))

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
